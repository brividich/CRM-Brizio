"""MOD.128 MPQ — F3 export "replica-Word" (.docx) della vista MOD.128.

Costruisce un documento Word fedele al modulo cartaceo MOD.128: una **tabella per
cliente** con una riga per processo qualificato e le colonne del modulo
(personale qualificato/addetto/controllore/part145, scadenze, distribuzione a
reparto). Funzione **pura** (nessun accesso al DB): riceve la struttura già
aggregata da ``views_mpq._build_vista_groups`` e ritorna i **byte** del .docx —
così è testabile riaprendo il documento con python-docx.

Dipendenza: ``python-docx`` (presente in ``requirements``).
"""
from __future__ import annotations

from datetime import date
from io import BytesIO

# Colonne della tabella (il CLIENTE è l'intestazione di sezione, come nel modulo).
_COLONNE = [
    "Processo qualificato",
    "Personale qualificato",
    "Addetto",
    "Controllore",
    "Part 145",
    "Scadenze",
    "Distribuzione a reparto",
]


def _set_cell_lines(cell, lines) -> None:
    """Scrive più righe in una cella (un paragrafo per voce); '—' se vuota."""
    valori = [str(x) for x in (lines or []) if str(x).strip()]
    cell.text = valori[0] if valori else "—"
    for extra in valori[1:]:
        cell.add_paragraph(extra)


def _riga_processo_cell_processo(riga) -> str:
    """Prima colonna: nome processo + livello/regime + riferimenti (compatti)."""
    testa = str(riga.get("processo") or "").strip()
    meta = []
    if riga.get("regime"):
        meta.append(str(riga["regime"]))
    if riga.get("livello"):
        meta.append(f"Liv. {riga['livello']}")
    if riga.get("stato"):
        meta.append(str(riga["stato"]))
    parti = [testa]
    if meta:
        parti.append(" · ".join(meta))
    rif = [str(r) for r in (riga.get("riferimenti") or []) if str(r).strip()]
    if rif:
        parti.append("Rif.: " + ", ".join(rif))
    return "\n".join(parti)


def _personale_lines(riga, chiave):
    """Righe della colonna personale, o il rimando organizzativo se applicabile."""
    if riga.get("organizzativo"):
        # Il personale organizzativo non è nominale: rimando alla dichiarazione,
        # mostrato una sola volta nella colonna "Personale qualificato".
        return [riga.get("organizzativo_rif") or "Organizzativo"] if chiave == "qualificato" else []
    return riga.get(chiave) or []


def build_mod128_docx_bytes(gruppi, oggi: date | None = None) -> bytes:
    """Costruisce il .docx della vista MOD.128 e ritorna i byte.

    ``gruppi`` = lista di ``{"cliente": str, "righe": [dict, ...]}`` dove ogni
    riga espone processo/livello/regime/riferimenti, le liste personale per
    ruolo (``qualificato``/``addetto``/``controllore``/``part145``), ``scadenze``,
    ``reparti``, ``stato`` e i flag ``organizzativo``/``organizzativo_rif``.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    titolo = doc.add_heading("MOD.128 — Mansionario Processi Qualificati (MPQ)", level=1)
    for run in titolo.runs:
        run.font.size = Pt(15)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        f"Vista processi qualificati · estratto del {oggi:%d-%m-%Y}" if oggi
        else "Vista processi qualificati"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(9)

    if not gruppi:
        doc.add_paragraph("Nessun processo qualificato registrato.")
    for gruppo in gruppi:
        doc.add_heading(f"Cliente/ente: {gruppo.get('cliente') or '—'}", level=2)
        righe = gruppo.get("righe") or []
        table = doc.add_table(rows=1, cols=len(_COLONNE))
        try:
            table.style = "Table Grid"
        except KeyError:  # pragma: no cover - stile sempre presente nel template base
            pass
        hdr = table.rows[0].cells
        for i, nome in enumerate(_COLONNE):
            hdr[i].text = nome
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(8)
        for riga in righe:
            cells = table.add_row().cells
            cells[0].text = _riga_processo_cell_processo(riga)
            _set_cell_lines(cells[1], _personale_lines(riga, "qualificato"))
            _set_cell_lines(cells[2], _personale_lines(riga, "addetto"))
            _set_cell_lines(cells[3], _personale_lines(riga, "controllore"))
            _set_cell_lines(cells[4], _personale_lines(riga, "part145"))
            _set_cell_lines(cells[5], riga.get("scadenze") or [])
            _set_cell_lines(cells[6], riga.get("reparti") or [])
            for c in cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
